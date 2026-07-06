from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from openpyxl import load_workbook

from everida.schemas import DocumentSection, DocumentTable, ParsedDocument, SpreadsheetSheet


def parse_document(path: str | Path) -> ParsedDocument:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return parse_docx(file_path)
    if suffix == ".xlsx":
        return parse_xlsx(file_path)
    if suffix == ".pdf":
        return ParsedDocument(
            kind="pdf",
            path=str(file_path),
            markdown="PDF parsing is reserved for the Docling/MinerU integration.",
        )
    return parse_text(file_path)


def parse_docx(path: Path) -> ParsedDocument:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    sections = _sections_from_paragraphs(paragraphs)
    tables: list[DocumentTable] = []

    markdown_parts: list[str] = []
    for index, paragraph in enumerate(paragraphs, start=1):
        markdown_parts.append(paragraph)
        if _looks_like_heading(paragraph):
            markdown_parts[-1] = f"## {paragraph}"

    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append(DocumentTable(rows=rows, source_ref=f"docx:table:{table_index}"))
            markdown_parts.append(_table_to_markdown(rows))

    return ParsedDocument(
        kind="docx",
        path=str(path),
        markdown="\n\n".join(part for part in markdown_parts if part),
        sections=sections,
        tables=tables,
    )


def parse_xlsx(path: Path) -> ParsedDocument:
    workbook = load_workbook(path, data_only=True)
    sheets: list[SpreadsheetSheet] = []
    markdown_parts: list[str] = []

    for worksheet in workbook.worksheets:
        rows = [
            [cell for cell in row]
            for row in worksheet.iter_rows(values_only=True)
            if any(cell is not None for cell in row)
        ]
        normalized = [[_normalize_cell(cell) for cell in row] for row in rows]
        sheets.append(
            SpreadsheetSheet(
                name=worksheet.title,
                rows=normalized,
                source_ref=f"xlsx:sheet:{worksheet.title}",
            )
        )
        markdown_parts.append(f"## {worksheet.title}\n\n{_table_to_markdown(normalized[:30])}")

    return ParsedDocument(
        kind="xlsx",
        path=str(path),
        markdown="\n\n".join(markdown_parts),
        sheets=sheets,
    )


def parse_text(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8")
    return ParsedDocument(
        kind="text",
        path=str(path),
        markdown=text,
        sections=[DocumentSection(title=path.name, text=text, source_ref=f"text:{path.name}")],
    )


def _sections_from_paragraphs(paragraphs: Iterable[str]) -> list[DocumentSection]:
    sections: list[DocumentSection] = []
    current_title = "正文"
    current_lines: list[str] = []
    current_ref = "docx:paragraph:1"

    for index, paragraph in enumerate(paragraphs, start=1):
        if _looks_like_heading(paragraph):
            if current_lines:
                sections.append(
                    DocumentSection(
                        title=current_title,
                        text="\n".join(current_lines),
                        source_ref=current_ref,
                    )
                )
            current_title = paragraph
            current_lines = []
            current_ref = f"docx:paragraph:{index}"
        else:
            current_lines.append(paragraph)

    if current_lines or not sections:
        sections.append(
            DocumentSection(
                title=current_title,
                text="\n".join(current_lines),
                source_ref=current_ref,
            )
        )
    return sections


def _looks_like_heading(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    heading_prefixes = ("第", "一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、")
    return len(stripped) <= 40 and (
        stripped.startswith(heading_prefixes)
        or stripped.endswith("：")
        or stripped.endswith(":")
        or stripped[0].isdigit()
    )


def _table_to_markdown(rows: list[list[object]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [[_normalize_cell(cell) for cell in row] + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]
    markdown_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(str(cell) for cell in row) + " |" for row in markdown_rows)


def _normalize_cell(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()
