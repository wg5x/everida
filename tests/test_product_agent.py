import json
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from everida.agents.product import (
    fill_template,
    generate_sql,
    parse_product,
    run_product_pipeline,
    validate_product,
)


ROOT = Path(__file__).resolve().parents[1]
SPEC = ROOT / "raw" / "【已评审需规】120078-未来星年金保险（分红.docx"
TEMPLATE = ROOT / "raw" / "产品配置模板.xlsx"
SQL = ROOT / "raw" / "120078_N_1_1.sql"


def test_parse_product_extracts_core_120078_fields():
    product = parse_product(SPEC)

    assert product.risk_code == "120078"
    assert "未来星" in product.risk_name
    assert "保障方案一" in product.coverage_schemes
    assert "保障方案二" in product.coverage_schemes
    assert product.product_type == "年金保险"
    assert "分红型" == product.bonus_type
    assert "趸交" in product.payment_options
    assert "3年" in product.payment_options
    assert "至25岁" in product.insurance_periods
    assert "成长守护金" in product.benefit_rules


def test_parse_product_extracts_code_and_name_without_120078_hardcoding(tmp_path):
    spec = _write_custom_product_docx(tmp_path)

    product = parse_product(spec)

    assert product.risk_code == "999001"
    assert product.risk_name == "星河年金保险（分红型）"
    assert product.short_name == "星河"
    assert product.coverage_schemes == ["保障方案A", "保障方案B"]
    assert "120078" not in product.model_dump_json()


def test_parse_product_includes_field_level_evidence():
    product = parse_product(SPEC)

    risk_code_evidence = product.field_evidence["risk_code"]
    assert risk_code_evidence.value == "120078"
    assert risk_code_evidence.source_ref.startswith("docx:")
    assert risk_code_evidence.confidence >= 0.8
    assert "120078" in risk_code_evidence.evidence_text

    payment_evidence = product.field_evidence["payment_options"]
    assert "3年" in payment_evidence.value
    assert payment_evidence.confidence >= 0.6

    benefit_evidence = product.field_evidence["benefit_rules"]
    assert "成长守护金" in benefit_evidence.value
    assert "成长守护金" in benefit_evidence.evidence_text


def test_parse_product_prioritizes_structured_context_for_payment_evidence():
    product = parse_product(SPEC)

    payment_evidence = product.field_evidence["payment_options"]
    assert payment_evidence.source_ref.startswith("docx:table:")
    assert "缴费期间" in payment_evidence.evidence_text or "缴费频率" in payment_evidence.evidence_text
    assert "3-Y-年" in payment_evidence.evidence_text


def test_fill_template_writes_product_summary_sheet(tmp_path):
    product = parse_product(SPEC)
    output = tmp_path / "filled.xlsx"

    fill_template(product, TEMPLATE, output)

    workbook = load_workbook(output)
    assert "Everida产品摘要" in workbook.sheetnames
    sheet = workbook["Everida产品摘要"]
    assert sheet["A1"].value == "字段"
    assert sheet["B2"].value == "120078"
    assert "coverage_schemes" in [row[0].value for row in sheet.iter_rows()]


def test_fill_template_maps_product_base_info_into_original_sheet(tmp_path):
    product = parse_product(SPEC)
    output = tmp_path / "filled.xlsx"

    fill_template(product, TEMPLATE, output)

    workbook = load_workbook(output)
    sheet = workbook["产品基础信息"]
    assert sheet["D2"].value == "120078"
    assert sheet["D3"].value == "中邮未来星年金保险（分红型）"
    assert sheet["D4"].value == "未来星"
    assert sheet["D2"].style_id == load_workbook(TEMPLATE)["产品基础信息"]["D2"].style_id


def test_fill_template_maps_duty_definitions_into_original_sheet(tmp_path):
    spec = _write_custom_product_docx(tmp_path)
    product = parse_product(spec)
    output = tmp_path / "filled.xlsx"

    fill_template(product, TEMPLATE, output)

    workbook = load_workbook(output)
    original = load_workbook(TEMPLATE)
    sheet = workbook["责任定义"]
    original_sheet = original["责任定义"]
    assert sheet["B2"].value == "保障方案A责任"
    assert sheet["B3"].value == "保障方案B责任"
    assert sheet["C2"].value == "10"
    assert sheet["E2"].value == "30"
    assert sheet["B2"].style_id == original_sheet["B2"].style_id
    assert sheet["C2"].style_id == original_sheet["C2"].style_id
    assert sheet["E2"].style_id == original_sheet["E2"].style_id


def test_fill_template_maps_benefit_rules_into_original_sheet(tmp_path):
    spec = _write_custom_product_docx(tmp_path)
    product = parse_product(spec)
    output = tmp_path / "filled.xlsx"

    fill_template(product, TEMPLATE, output)

    workbook = load_workbook(output)
    original = load_workbook(TEMPLATE)
    sheet = workbook["给付责任"]
    original_sheet = original["给付责任"]
    assert [sheet[f"C{row}"].value for row in range(2, 5)] == [
        "成长守护金",
        "满期保险金",
        "身故保险金",
    ]
    assert sheet["C2"].style_id == original_sheet["C2"].style_id
    assert sheet["C3"].style_id == original_sheet["C3"].style_id
    assert sheet["C4"].style_id == original_sheet["C4"].style_id


def test_fill_template_preserves_original_workbook_structure_and_styles(tmp_path):
    product = parse_product(SPEC)
    output = tmp_path / "filled.xlsx"
    original = load_workbook(TEMPLATE)
    original_sheet_names = list(original.sheetnames)
    original_snapshot = _workbook_template_snapshot(
        original,
        ignored_cells={"产品基础信息": {"D2", "D3", "D4"}, "给付责任": {f"C{row}" for row in range(2, 10)}},
    )

    fill_template(product, TEMPLATE, output)

    filled = load_workbook(output)
    assert filled.sheetnames[: len(original_sheet_names)] == original_sheet_names
    assert filled.sheetnames[-1] == "Everida产品摘要"
    assert _workbook_template_snapshot(
        filled,
        original_sheet_names,
        ignored_cells={"产品基础信息": {"D2", "D3", "D4"}, "给付责任": {f"C{row}" for row in range(2, 10)}},
    ) == original_snapshot


def test_generate_sql_contains_draft_guard_and_core_fields():
    product = parse_product(SPEC)
    sql = generate_sql(product)

    assert "Everida generated draft SQL" in sql
    assert "120078" in sql
    assert "成长守护金" in sql


def test_validate_product_outputs_markdown_report():
    product = parse_product(SPEC)
    report = validate_product(product, TEMPLATE, SQL)

    assert "# Everida 产品一致性校验报告" in report
    assert "120078" in report
    assert "| 保障方案 |" in report
    assert "字段证据链" in report
    assert "人工确认项" in report


def test_run_product_pipeline_writes_all_mvp_artifacts(tmp_path):
    result = run_product_pipeline(SPEC, TEMPLATE, SQL, tmp_path)

    expected = {
        "parsed_document_json",
        "parsed_document_markdown",
        "product_json",
        "filled_xlsx",
        "generated_sql",
        "sql_inventory_json",
        "validate_report_markdown",
        "run_manifest_json",
    }
    assert expected == set(result.artifacts)
    for path in result.artifacts.values():
        assert Path(path).exists()

    product = json.loads((tmp_path / "product.json").read_text(encoding="utf-8"))
    manifest = json.loads((tmp_path / "run_manifest.json").read_text(encoding="utf-8"))
    assert product["risk_code"] == "120078"
    assert product["field_evidence"]["risk_code"]["value"] == "120078"
    assert manifest["status"] == "completed"


def _workbook_template_snapshot(workbook, sheet_names: list[str] | None = None, ignored_cells=None):
    names = sheet_names or list(workbook.sheetnames)
    ignored_cells = ignored_cells or {}
    snapshot = {}
    for sheet_name in names:
        sheet = workbook[sheet_name]
        column_widths = {
            key: dimension.width
            for key, dimension in sheet.column_dimensions.items()
            if dimension.width is not None
        }
        row_heights = {
            key: dimension.height
            for key, dimension in sheet.row_dimensions.items()
            if dimension.height is not None
        }
        cells = {}
        for row in sheet.iter_rows():
            for cell in row:
                if cell.coordinate in ignored_cells.get(sheet_name, set()):
                    continue
                if cell.value is not None:
                    cells[cell.coordinate] = {
                        "value": cell.value,
                        "style_id": cell.style_id,
                        "comment": cell.comment.text if cell.comment else None,
                    }
        snapshot[sheet_name] = {
            "max_row": sheet.max_row,
            "max_column": sheet.max_column,
            "merged_ranges": sorted(str(cell_range) for cell_range in sheet.merged_cells.ranges),
            "column_widths": column_widths,
            "row_heights": row_heights,
            "cells": cells,
        }
    return snapshot


def _write_custom_product_docx(tmp_path: Path) -> Path:
    spec = tmp_path / "custom_product.docx"
    document = Document()
    document.add_paragraph("999001-星河年金保险（分红型）")
    document.add_paragraph("产品类型：年金保险")
    document.add_paragraph("保障方案：保障方案A、保障方案B")
    table = document.add_table(rows=3, cols=4)
    table.cell(0, 0).text = "缴费期间*"
    table.cell(0, 1).text = "3-Y-年"
    table.cell(0, 2).text = "5-Y-年"
    table.cell(0, 3).text = "10-Y-年"
    table.cell(1, 0).text = "保险期间*"
    table.cell(1, 1).text = "至25岁"
    table.cell(1, 2).text = "至30岁"
    table.cell(1, 3).text = "30-Y-年"
    table.cell(2, 0).text = "给付责任名称*"
    table.cell(2, 1).text = "成长守护金"
    table.cell(2, 2).text = "满期保险金"
    table.cell(2, 3).text = "身故保险金"
    document.save(spec)
    return spec
